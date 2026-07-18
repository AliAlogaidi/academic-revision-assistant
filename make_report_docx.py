from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SOURCE = Path("Doctor_Submission_Report.md")
OUTPUT = Path("AI_Academic_Revision_Assistant_Report.docx")


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value


def main() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0

    while index < len(lines):
        line = lines[index].strip()

        if not line or line == "---":
            index += 1
            continue

        if line.startswith("|"):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                row = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(set(cell) <= {"-"} for cell in row):
                    rows.append(row)
                index += 1
            if rows:
                add_table(doc, rows)
            continue

        if line.startswith("```"):
            block: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            paragraph = doc.add_paragraph("\n".join(block))
            paragraph.style = "No Spacing"
        elif line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(16)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("**") and ":**" in line:
            doc.add_paragraph(line.replace("**", ""))
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

        index += 1

    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
